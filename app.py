import streamlit as st
import docx
from styling import apply_css_styling
from summify_functions import split_files_into_groups, generate_final_summary
import time
import openai
apply_css_styling()


# OpenAI API configuration
openai.api_type = 'azure'
openai.api_version = '2022-12-01'
deployment_name = 'ishita-davinci-test-3'

openai.api_key = "39cdcddc7d15446c8588a34eada4d121"
openai.api_base = "https://testnew.openai.azure.com/"

col1, col2 = st.columns([2, 1])

# First column content
col1.markdown("<div class='title-container' style='padding-top:140px; padding-bottom: 40px;'><span class='title-text'>Summify-it</span></div>", unsafe_allow_html=True)
col1.markdown("<div class='subtitle-container'><span class='subtitle-text'>An AI-assisted Summary Tool for your Files</span></div>", unsafe_allow_html=True)

# Second column content (contains image)
col2.image('image.png')

# Position the expander in the top right corner
st.sidebar.markdown("<div class='sidebar-header-container'><span class='sidebar-header-text'>Test Guide</span></div>", unsafe_allow_html=True)

st.divider()
with st.container():
    st.markdown("<div class='upload-header-container'><span class='upload-header-text'>Upload Files: </span></div>", unsafe_allow_html=True)
    uploaded_files = st.file_uploader(label='---------------------------------', accept_multiple_files=True)
    
if uploaded_files:
        st.success("Documents uploaded successfully!")
        
        if st.button('Summarize'):
            start_time = time.time()  # Start measuring the execution time
            
            
            text_to_summarize_list = []
            try:
                file_groups = split_files_into_groups(uploaded_files, group_size=3)
                for files in file_groups:
                    text_to_summarize = ""
                    for uploaded_file in files:
                        doc = docx.Document(uploaded_file)
                        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                        text_to_summarize += text
                        for table in doc.tables:
                            table_text = ""
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        table_text += paragraph.text + " "
                            text_to_summarize += table_text
                    text_to_summarize_list.append(text_to_summarize)
                st.divider()
                with st.spinner('Summarizing documents...'):
                    final_summary = generate_final_summary(text_to_summarize_list)
                st.success("Documents summarized successfully!")
                font_family = "Arial"
                font_size = "16px"
                font_style = "normal"
                final_summary_with_font = f"<span style='font-family: {font_family}; font-size: {font_size}; font-style: {font_style};'>{final_summary}</span>"
                st.markdown(
                    
                f"""
                <div style='border: 1px solid black; padding: 10px; border-radius: 10px; margin: 10px;'>
                   <h2>Summary</h2> {final_summary_with_font}
                </div>
                """,
                unsafe_allow_html=True
                )
                end_time = time.time()  # End measuring the execution time
                execution_time = end_time - start_time
                st.sidebar.info(f"Execution Time for Query: {execution_time:.2f} seconds")
                st.download_button('Download Summary', final_summary, file_name='summary.txt', key="download-button")
            except Exception as e:
                st.error(f"Error occurred: {str(e)}")
        else:
            st.sidebar.info("Get your query execution time here !")

with st.sidebar:
    st.markdown(
        """
        <style>
        .expander-content 
        {
        padding: 5px; 
        border-radius: 20px; 
        } </style> """, 
        unsafe_allow_html=True,
    )

    app_usage=["1. Click on browse files on the sidebar.","2. Upload files you wish to summarize.","3. Click on the Summarize button.","4. Now you can download the summary."]
    with st.sidebar.expander("How to Use?", expanded=False):
        if app_usage:
            st.markdown("<ol class='helper-prompts'>", unsafe_allow_html=True)
            for uses in app_usage:
                st.markdown(f"<li style='margin-left:-10px; font-size:13px;'>{uses}</li>", unsafe_allow_html=True)
            st.markdown("</ol>", unsafe_allow_html=True)

vidfile_path="usage.mp4"
st.sidebar.video(vidfile_path, format='video/mp4', start_time=0)